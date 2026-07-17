import sys
import os
from pathlib import Path
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QProgressBar, QTextEdit,
    QFileDialog, QMessageBox
)
from PyQt5.QtCore import QThread, pyqtSignal, Qt
from PyQt5.QtGui import QFont
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
 
class Worker(QThread):
    progress = pyqtSignal(int)
    log_message = pyqtSignal(str)
    finished = pyqtSignal(bool, str)
 
    def __init__(self, input_path: str, output_path: str):
        super().__init__()
        self.input_path = input_path
        self.output_path = output_path
        self.debug_log_path = Path(output_path).parent / (Path(output_path).stem + "_debug.log")
 
    def run(self):
        try:
            doc = Document(self.input_path)
            total = self._count_paragraphs(doc)
            processed = 0
            changed_count = 0
 
            with open(self.debug_log_path, 'w', encoding='utf-8') as debug_file:
                debug_file.write("Bold Normalization Debug Log (XML edition)\n")
                debug_file.write(f"Input:  {self.input_path}\n")
                debug_file.write(f"Output: {self.output_path}\n\n")
 
                self.log_message.emit(f"Processing {total} paragraphs using XML editing...")
 
                for para in self._iterate_paragraphs(doc):
                    processed += 1
                    action, details = self._process_paragraph_xml(para)
                    if action == 'changed':
                        changed_count += 1
                        debug_file.write(f"Paragraph {processed}: CHANGED\n")
                        debug_file.write(f"  Text: {para.text[:80]}...\n")
                        debug_file.write(f"  Details: {details}\n\n")
                    else:
                        debug_file.write(f"Paragraph {processed}: {action}\n")
                        debug_file.write(f"  Text: {para.text[:80]}...\n\n")
 
                    progress_val = int(processed / total * 100)
                    self.progress.emit(progress_val)
 
                    if processed % 50 == 0:
                        self.log_message.emit(f"Processed {processed} of {total} paragraphs...")
 
                doc.save(self.output_path)
 
            self.log_message.emit(f"Done. Changed {changed_count} paragraphs.")
            self.finished.emit(True, f"Saved to {self.output_path}\nDebug log: {self.debug_log_path}")
 
        except Exception as e:
            self.log_message.emit(f"ERROR: {str(e)}")
            self.finished.emit(False, str(e))
 
    def _count_paragraphs(self, doc) -> int:
        count = 0
        for _ in self._iterate_paragraphs(doc):
            count += 1
        return count
 
    def _iterate_paragraphs(self, doc):
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in doc.sections:
            yield from section.header.paragraphs
            yield from section.footer.paragraphs
 
    def _is_run_bold(self, run, para):
        """Return True if this run is effectively bold (direct or inherited).
 
        NOTE: python-docx's Run object has no public `.parent` attribute,
        so the paragraph must be passed in explicitly (that was the bug).
        """
        # 1. Check direct <w:b> or <w:bCs> on the run element
        r = run._element
        for tag in (qn('w:b'), qn('w:bCs')):
            elem = r.find(tag)
            if elem is not None:
                # A bold element with w:val="0"/"false" explicitly disables bold
                val = elem.get(qn('w:val'))
                if val is not None and val.lower() in ('0', 'false', 'off'):
                    return False
                return True
 
        # 2. Check run.bold property (may be True/False/None)
        if run.bold is not None:
            return run.bold
 
        # 3. Inherit from paragraph style
        if para is not None and para.style is not None:
            style = para.style
            if style.font and style.font.bold is not None:
                return style.font.bold
 
        # 4. Inherit from character style (if any)
        if run.style is not None:
            if run.style.font and run.style.font.bold is not None:
                return run.style.font.bold
 
        # 5. Default: not bold
        return False
 
    def _remove_bold_from_run(self, run):
        """Force-disable bold on this run's XML, for BOTH the regular
        (w:b) and complex-script (w:bCs) properties.
 
        This matters a lot for Arabic/RTL text: Word renders bold for
        complex-script runs based on <w:bCs>, not <w:b>. python-docx's
        `run.bold = False` setter only ever writes <w:b>, so relying on
        it alone leaves <w:bCs> absent -> Word falls back to whatever
        the paragraph/character style says for complex-script bold,
        which is frequently still "true". That's why runs could still
        render bold even though they were "fixed" here.
 
        To guarantee an override regardless of style inheritance, we
        explicitly set w:val="0" on both elements directly in the XML.
        """
        r = run._element
        rPr = r.get_or_add_rPr()
        for tag in ('w:b', 'w:bCs'):
            elem = rPr.find(qn(tag))
            if elem is None:
                elem = OxmlElement(tag)
                rPr.append(elem)
            elem.set(qn('w:val'), '0')
 
    def _process_paragraph_xml(self, para):
        """
        Inspect each run's bold status.
        If mixed (some bold, some not), remove bold from all runs.
        Returns (action, details).
        """
        runs = para.runs
        if not runs:
            return 'no_runs', ''
 
        # Determine bold status for each run
        statuses = []
        for run in runs:
            statuses.append(self._is_run_bold(run, para))
 
        # Check if all are bold or all are regular
        all_bold = all(statuses)
        all_regular = not any(statuses)
 
        if all_bold:
            return 'all_bold', ''
        if all_regular:
            return 'all_regular', ''
 
        # Mixed: remove bold from every run
        for run in runs:
            self._remove_bold_from_run(run)
 
        details = f"Runs: {len(runs)}, statuses: {statuses}"
        return 'changed', details
 
 
# -------------------- GUI (unchanged) --------------------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Bold Normalizer (XML) – Arabic Documents")
        self.setGeometry(100, 100, 750, 600)
 
        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)
 
        input_layout = QHBoxLayout()
        input_layout.addWidget(QLabel("Input:"))
        self.input_edit = QLineEdit()
        self.input_browse = QPushButton("Browse…")
        input_layout.addWidget(self.input_edit)
        input_layout.addWidget(self.input_browse)
        layout.addLayout(input_layout)
 
        output_layout = QHBoxLayout()
        output_layout.addWidget(QLabel("Output:"))
        self.output_edit = QLineEdit()
        self.output_browse = QPushButton("Browse…")
        output_layout.addWidget(self.output_edit)
        output_layout.addWidget(self.output_browse)
        layout.addLayout(output_layout)
 
        self.start_btn = QPushButton("Start Processing")
        layout.addWidget(self.start_btn)
 
        self.progress = QProgressBar()
        layout.addWidget(self.progress)
 
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setFont(QFont("Consolas", 9))
        layout.addWidget(self.log_text)
 
        self.input_browse.clicked.connect(self.browse_input)
        self.output_browse.clicked.connect(self.browse_output)
        self.start_btn.clicked.connect(self.start_processing)
 
    def browse_input(self):
        file, _ = QFileDialog.getOpenFileName(
            self, "Select Word Document", "", "Word Documents (*.docx)"
        )
        if file:
            self.input_edit.setText(file)
            if not self.output_edit.text():
                path = Path(file)
                out_path = path.parent / (path.stem + "_fixed.docx")
                self.output_edit.setText(str(out_path))
 
    def browse_output(self):
        file, _ = QFileDialog.getSaveFileName(
            self, "Save Fixed Document", "", "Word Documents (*.docx)"
        )
        if file:
            self.output_edit.setText(file)
 
    def start_processing(self):
        input_path = self.input_edit.text().strip()
        output_path = self.output_edit.text().strip()
 
        if not input_path or not os.path.exists(input_path):
            QMessageBox.warning(self, "Error", "Please select a valid input file.")
            return
        if not output_path:
            QMessageBox.warning(self, "Error", "Please specify an output file path.")
            return
 
        self.start_btn.setEnabled(False)
        self.progress.setValue(0)
        self.log_text.clear()
        self.log_text.append("Starting processing (XML editing)...")
 
        self.worker = Worker(input_path, output_path)
        self.worker.progress.connect(self.update_progress)
        self.worker.log_message.connect(self.append_log)
        self.worker.finished.connect(self.on_finished)
        self.worker.start()
 
    def update_progress(self, value):
        self.progress.setValue(value)
 
    def append_log(self, msg):
        self.log_text.append(msg)
 
    def on_finished(self, success, message):
        self.start_btn.setEnabled(True)
        if success:
            QMessageBox.information(self, "Success", message)
        else:
            QMessageBox.critical(self, "Error", f"Processing failed:\n{message}")
        self.append_log(message)
 
 
if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
 