import argparse
import os
import sys
import tempfile
import time
from pathlib import Path

try:
    from PyQt5 import QtWidgets
except Exception as exc:  # noqa: BLE001
    print("PyQt5 is required to run this application.")
    print(str(exc))
    sys.exit(1)

from core.engines import DocxEngine, EngineUnavailableError, get_runtime_status
from core.operations import Operation, OperationRegistry, OperationResult
from core.pipeline import Pipeline
import operations  # noqa: F401
from ui.main_window import MainWindow


APP_TITLE = "Batch Edit Files"
PROJECT_ROOT = Path(__file__).resolve().parent


def find_sample_doc() -> Path:
    candidates = [
        PROJECT_ROOT / "clean.docx",
        PROJECT_ROOT / "_smoke" / "test_sections.docx",
    ]
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    raise FileNotFoundError("No sample DOCX file found for smoke testing.")


def create_phase4_sample_doc(
    output_path: Path,
    body_lines: list[str],
    header_text: str = "",
) -> None:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise EngineUnavailableError(
            "python-docx is required to build the Phase 4 smoke sample."
        ) from exc

    doc = Document()
    for line in body_lines:
        doc.add_paragraph(line)
    if header_text:
        section = doc.sections[0]
        section.header.paragraphs[0].text = header_text
    doc.save(str(output_path))


def run_phase1_smoke() -> int:
    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)
    window = MainWindow(app_title=APP_TITLE)
    assert window.windowTitle() == APP_TITLE
    assert window.main_tabs.count() == 2
    assert window.statusBar() is not None
    print("Phase 1 smoke test passed: app shell launched with two tabs and status bar.")
    window.close()
    app.quit()
    return 0


def run_phase2_smoke() -> int:
    sample_doc = find_sample_doc()
    runtime_status = get_runtime_status()
    lines = [f"Sample DOCX: {sample_doc}"]

    if runtime_status.docx_available:
        engine = DocxEngine()
        try:
            engine.open(str(sample_doc))
            styles = engine.list_styles()
            paragraphs = engine.get_paragraphs()
            lines.append(
                f"python-docx engine opened document, found {len(styles)} styles and "
                f"{len(paragraphs)} paragraphs."
            )
        finally:
            engine.close()
    else:
        lines.append(f"python-docx engine unavailable: {runtime_status.docx_reason}")

    if runtime_status.com_available:
        com_engine = None
        try:
            com_engine = runtime_status.create_com_engine()
            com_engine.open(str(sample_doc))
            com_styles = com_engine.list_styles()
            lines.append(f"COM engine opened document, found {len(com_styles)} styles.")
        except Exception as exc:  # noqa: BLE001
            lines.append(f"COM engine probe failed during smoke test: {exc}")
            return_code = 1
        else:
            return_code = 0
        finally:
            if com_engine is not None:
                com_engine.close()
    else:
        lines.append(f"COM engine unavailable: {runtime_status.com_reason}")
        return_code = 0

    print("\n".join(lines))
    return return_code


def run_phase3_smoke() -> int:
    class MockEngine:
        def __init__(self) -> None:
            self.capabilities = {"smoke"}
            self.events = []

    if "Smoke Append" not in OperationRegistry.list_operations():
        @OperationRegistry.register("Smoke Append")
        class SmokeAppendOperation(Operation):
            name = "Smoke Append"
            required_capabilities = {"smoke"}

            def validate(self, engine) -> None:
                self.ensure_supported(engine)

            def run(self, engine) -> OperationResult:
                engine.events.append(self.params["value"])
                return OperationResult(
                    status="ok",
                    message=f"Appended {self.params['value']}",
                    details={"events": list(engine.events)},
                )

            def describe(self) -> str:
                return f"Append {self.params['value']}"

    class SmokeFailureOperation(Operation):
        name = "Smoke Failure"
        required_capabilities = {"smoke"}

        def validate(self, engine) -> None:
            self.ensure_supported(engine)

        def run(self, engine) -> OperationResult:
            return OperationResult(status="error", message="Intentional smoke failure")

        def describe(self) -> str:
            return "Fail intentionally"

    engine = MockEngine()
    pipeline = Pipeline(
        operations=[
            OperationRegistry.create("Smoke Append", {"value": "first"}),
            OperationRegistry.create("Smoke Append", {"value": "second"}),
            SmokeFailureOperation(),
        ],
        continue_on_error=True,
    )
    results = pipeline.run(engine)

    assert len(results) == 3
    assert [result.status for result in results] == ["ok", "ok", "error"]
    assert engine.events == ["first", "second"]
    print("Phase 3 smoke test passed: pipeline ordering and error propagation verified.")
    return 0


def run_phase4_smoke() -> int:
    from docx import Document

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase4_") as temp_dir:
        temp_root = Path(temp_dir)

        docx_sample = temp_root / "phase4_docx.docx"
        create_phase4_sample_doc(
            docx_sample,
            body_lines=["alpha beta alpha", "alphax alphay"],
        )

        docx_engine = DocxEngine()
        docx_engine.open(str(docx_sample))
        docx_result = OperationRegistry.create(
            "Find/Replace",
            {
                "find_text": "alpha",
                "replace_text": "omega",
                "scope": {"body": True},
                "match_case": False,
                "whole_word": True,
            },
        ).run(docx_engine)
        docx_engine.save()
        docx_engine.close()

        docx_doc = Document(str(docx_sample))
        docx_body = [paragraph.text for paragraph in docx_doc.paragraphs]
        assert docx_result.details["count"] == 2
        assert docx_body[0] == "omega beta omega"
        assert docx_body[1] == "alphax alphay"

        runtime_status = get_runtime_status()
        lines = [
            "Phase 4 smoke test passed for python-docx exact replace.",
            f"python-docx result count: {docx_result.details['count']}",
        ]

        if runtime_status.com_available:
            com_sample = temp_root / "phase4_com.docx"
            create_phase4_sample_doc(
                com_sample,
                body_lines=["alpha beta alpha", "alphax alphay"],
                header_text="header alpha",
            )
            com_engine = runtime_status.create_com_engine()
            try:
                com_engine.open(str(com_sample))
                exact_result = OperationRegistry.create(
                    "Find/Replace",
                    {
                        "find_text": "alpha",
                        "replace_text": "omega",
                        "scope": {"body": True, "headers_footers": True},
                        "match_case": False,
                        "whole_word": True,
                    },
                ).run(com_engine)
                com_engine.save()
            finally:
                com_engine.close()

            com_doc = Document(str(com_sample))
            com_body = [paragraph.text for paragraph in com_doc.paragraphs]
            header_text = com_doc.sections[0].header.paragraphs[0].text
            assert exact_result.details["count"] == 3
            assert com_body[0] == "omega beta omega"
            assert com_body[1] == "alphax alphay"
            assert header_text == "header omega"

            regex_sample = temp_root / "phase4_regex.docx"
            create_phase4_sample_doc(
                regex_sample,
                body_lines=["alpha-123 alpha-456"],
            )
            regex_engine = runtime_status.create_com_engine()
            try:
                regex_engine.open(str(regex_sample))
                regex_result = OperationRegistry.create(
                    "Find/Replace",
                    {
                        "find_text": r"alpha-(\d+)",
                        "replace_text": r"id-\1",
                        "scope": {"body": True},
                        "use_regex": True,
                        "match_case": True,
                    },
                ).run(regex_engine)
                regex_engine.save()
            finally:
                regex_engine.close()

            regex_doc = Document(str(regex_sample))
            regex_body = [paragraph.text for paragraph in regex_doc.paragraphs]
            assert regex_result.details["count"] == 2
            assert regex_body[0] == "id-123 id-456"
            lines.append(
                "Phase 4 smoke test passed for COM exact + regex replace, including header scope."
            )
            lines.append(f"COM exact result count: {exact_result.details['count']}")
            lines.append(f"COM regex result count: {regex_result.details['count']}")
        else:
            lines.append(f"COM smoke skipped: {runtime_status.com_reason}")

        print("\n".join(lines))
        return 0


def create_phase5_sample_doc(
    output_path: Path,
    include_header: bool = True,
) -> None:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise EngineUnavailableError(
            "python-docx is required to build the Phase 5 smoke sample."
        ) from exc

    document = Document()
    document.add_paragraph("English paragraph should stay untouched.")
    document.add_paragraph("مرحبا بالعالم")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "العمود الأول"
    table.cell(0, 1).text = "العمود الثاني"
    if include_header:
        document.sections[0].header.paragraphs[0].text = "ترويسة عربية"
    document.save(str(output_path))


def run_phase5_smoke() -> int:
    from docx import Document

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase5_") as temp_dir:
        temp_root = Path(temp_dir)

        docx_sample = temp_root / "phase5_docx.docx"
        create_phase5_sample_doc(docx_sample, include_header=False)

        docx_engine = DocxEngine()
        docx_engine.open(str(docx_sample))
        docx_result = OperationRegistry.create(
            "Arabic RTL Normalize",
            {
                "scope": {"body": True},
                "alignment": "right",
                "direction": "rtl",
                "font_name": "Arial",
                "arabic_only": True,
                "normalize_tables": True,
            },
        ).run(docx_engine)
        docx_engine.save()
        docx_engine.close()

        docx_doc = Document(str(docx_sample))
        body_paragraph = docx_doc.paragraphs[1]
        body_ppr_xml = body_paragraph._p.get_or_add_pPr().xml
        run_xml = body_paragraph.runs[0]._r.get_or_add_rPr().xml
        table_xml = docx_doc.tables[0]._tbl.xml
        assert docx_result.details["paragraphs_updated"] >= 1
        assert "w:bidi" in body_ppr_xml
        assert 'w:val="right"' in body_ppr_xml
        assert "w:rtl" in run_xml
        assert "w:bidiVisual" in table_xml

        lines = [
            "Phase 5 smoke test passed for python-docx RTL normalization.",
            f"python-docx paragraphs updated: {docx_result.details['paragraphs_updated']}",
            f"python-docx tables updated: {docx_result.details['tables_updated']}",
        ]

        runtime_status = get_runtime_status()
        if runtime_status.com_available:
            com_sample = temp_root / "phase5_com.docx"
            create_phase5_sample_doc(com_sample, include_header=True)
            com_engine = runtime_status.create_com_engine()
            try:
                com_engine.open(str(com_sample))
                com_result = OperationRegistry.create(
                    "Arabic RTL Normalize",
                    {
                        "scope": {"body": True, "headers_footers": True},
                        "alignment": "right",
                        "direction": "rtl",
                        "font_name": "Arial",
                        "arabic_only": True,
                        "normalize_tables": True,
                    },
                ).run(com_engine)
                com_engine.save()
            finally:
                com_engine.close()

            com_doc = Document(str(com_sample))
            com_body = com_doc.paragraphs[1]
            com_body_xml = com_body._p.get_or_add_pPr().xml
            com_header_xml = com_doc.sections[0].header.paragraphs[0]._p.get_or_add_pPr().xml
            assert com_result.details["paragraphs_updated"] >= 2
            assert 'w:val="right"' in com_body_xml
            assert 'w:val="right"' in com_header_xml
            lines.append("Phase 5 smoke test passed for COM RTL normalization including header scope.")
            lines.append(f"COM paragraphs updated: {com_result.details['paragraphs_updated']}")
            lines.append(f"COM tables updated: {com_result.details['tables_updated']}")
        else:
            lines.append(f"COM smoke skipped: {runtime_status.com_reason}")

        print("\n".join(lines))
        return 0


def create_phase6_sample_doc(output_path: Path) -> None:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise EngineUnavailableError(
            "python-docx is required to build the Phase 6 smoke sample."
        ) from exc

    document = Document()
    document.add_paragraph("Document Title", style="Title")
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("Introductory paragraph for the first chapter.")
    document.add_heading("Section 1.1", level=2)
    document.add_paragraph("Supporting text for section 1.1.")
    document.add_heading("Chapter 2", level=1)
    document.add_paragraph("Introductory paragraph for the second chapter.")
    document.save(str(output_path))


def run_phase6_smoke() -> int:
    runtime_status = get_runtime_status()
    if not runtime_status.com_available:
        print(f"COM smoke skipped: {runtime_status.com_reason}")
        return 0

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase6_") as temp_dir:
        temp_root = Path(temp_dir)
        sample_path = temp_root / "phase6_toc.docx"
        create_phase6_sample_doc(sample_path)

        com_engine = runtime_status.create_com_engine()
        try:
            com_engine.open(str(sample_path))
            toc_result = OperationRegistry.create(
                "Custom TOC Builder",
                {
                    "style_levels": [
                        {"style_name": "Heading 1", "level": 1},
                        {"style_name": "Heading 2", "level": 2},
                    ],
                    "title_text": "Contents",
                    "tab_leader": "dots",
                    "show_page_numbers": True,
                    "use_hyperlinks": True,
                    "right_align_page_numbers": True,
                    "toc_font_name": "Arial",
                    "toc_font_size": 11,
                    "title_font_name": "Arial",
                    "title_font_size": 14,
                    "title_alignment": "center",
                    "insertion_location": "start",
                    "replace_existing": True,
                },
            ).run(com_engine)
            com_engine.save()
        finally:
            com_engine.close()

        verify_engine = runtime_status.create_com_engine()
        try:
            verify_engine.open(str(sample_path))
            document = verify_engine.document
            assert int(document.TablesOfContents.Count) == 1
            title_text = str(document.Paragraphs.Item(1).Range.Text).strip()
            toc = document.TablesOfContents.Item(1)
            try:
                toc.Update()
            except Exception:  # noqa: BLE001
                document.Fields.Update()
            toc_text = str(toc.Range.Text)
            assert title_text == "Contents"
            assert "Chapter 1" in toc_text
            assert "Section 1.1" in toc_text
            assert "Chapter 2" in toc_text
        finally:
            verify_engine.close()

        print("Phase 6 smoke test passed for COM TOC generation.")
        print(f"TOC count: {toc_result.details['toc_count']}")
        print(f"TOC preview contains headings: {'Chapter 1' in toc_result.details['toc_text_preview'] or 'Section 1.1' in toc_result.details['toc_text_preview']}")
        return 0


def run_phase7_smoke() -> int:
    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)
    window = MainWindow(app_title=APP_TITLE, show_summary_dialog=False)

    sample_doc = find_sample_doc()
    with tempfile.TemporaryDirectory(prefix="batch_edit_phase7_") as temp_dir:
        output_path = str(Path(temp_dir) / "phase7-output.docx")
        window.config_tab.default_output_row.set_text(temp_dir)
        window.load_document(str(sample_doc), output_path)

        assert window.current_engine is not None
        assert window.pipeline_tab.input_row.text() == str(sample_doc)
        assert window.pipeline_tab.output_row.text() == output_path
        assert window.pipeline_tab.styles_list.count() > 0
        assert window.pipeline_tab.operations_list.count() == len(
            OperationRegistry.list_operations()
        )
        assert "Active engine:" in window.pipeline_tab.active_engine_value.text()
        assert window.config_tab.active_engine_label.text().startswith("Active engine:")

        operation_rows = [
            window.pipeline_tab.operations_list.item(index).text()
            for index in range(window.pipeline_tab.operations_list.count())
        ]
        assert any("Find/Replace" in row for row in operation_rows)
        assert any("Arabic RTL Normalize" in row for row in operation_rows)
        assert any("Custom TOC Builder" in row for row in operation_rows)

        print(
            "Phase 7 smoke test passed: GUI workflow loaded a document, displayed the "
            "active engine, discovered styles, and evaluated operation availability."
        )

    window.close()
    app.quit()
    return 0


def run_phase8_smoke() -> int:
    from docx import Document

    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)
    window = MainWindow(app_title=APP_TITLE, show_summary_dialog=False)

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase8_") as temp_dir:
        temp_root = Path(temp_dir)
        input_path = temp_root / "phase8-input.docx"
        output_path = temp_root / "phase8-output.docx"
        create_phase5_sample_doc(input_path, include_header=False)

        window.config_tab.engine_preference_combo.setCurrentIndex(
            window.config_tab.engine_preference_combo.findData("docx")
        )
        window.config_tab.default_output_row.set_text(temp_dir)
        window.load_document(str(input_path), str(output_path))

        window.configured_operations = [
            {
                "display_name": "Find/Replace",
                "params": {
                    "find_text": "مرحبا",
                    "replace_text": "أهلا",
                    "scope": {"body": True},
                    "match_case": True,
                    "whole_word": False,
                },
                "last_status": "",
            },
            {
                "display_name": "Arabic RTL Normalize",
                "params": {
                    "scope": {"body": True},
                    "alignment": "right",
                    "direction": "rtl",
                    "font_name": "Arial",
                    "arabic_only": True,
                    "normalize_tables": True,
                    "normalize_lists": True,
                },
                "last_status": "",
            },
        ]
        window._refresh_configured_pipeline()
        assert window.pipeline_tab.configured_operations_list.count() == 2

        window.run_pipeline_from_ui()
        deadline = time.time() + 20
        while window.pipeline_worker is not None and time.time() < deadline:
            app.processEvents()
            time.sleep(0.05)
        assert window.pipeline_worker is None
        assert output_path.is_file()

        output_doc = Document(str(output_path))
        output_text = [paragraph.text for paragraph in output_doc.paragraphs]
        assert any("أهلا" in text for text in output_text)
        arabic_paragraph = next(
            paragraph for paragraph in output_doc.paragraphs if "أهلا" in paragraph.text
        )
        assert "w:bidi" in arabic_paragraph._p.get_or_add_pPr().xml
        assert window.pipeline_tab.log_panel.toPlainText().count("Running") >= 2

        print(
            "Phase 8 smoke test passed: GUI pipeline add/run flow saved output with "
            "operation results and logged per-operation progress."
        )

    window.close()
    app.quit()
    return 0


def run_phase9_smoke() -> int:
    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)
    window = MainWindow(app_title=APP_TITLE, show_summary_dialog=False)

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase9_") as temp_dir:
        temp_root = Path(temp_dir)
        input_path = temp_root / "phase9-input.docx"
        output_path = temp_root / "phase9-output.docx"
        log_path = temp_root / "phase9-session.jsonl"
        create_phase5_sample_doc(input_path, include_header=False)

        window.config_tab.engine_preference_combo.setCurrentIndex(
            window.config_tab.engine_preference_combo.findData("docx")
        )
        window.load_document(str(input_path), str(output_path))
        window.configured_operations = [
            {
                "display_name": "Find/Replace",
                "params": {
                    "find_text": "مرحبا",
                    "replace_text": "أهلا",
                    "scope": {"body": True},
                    "match_case": True,
                    "whole_word": False,
                },
                "last_status": "",
            },
            {
                "display_name": "Custom TOC Builder",
                "params": {
                    "style_levels": [{"style_name": "Heading 1", "level": 1}],
                    "title_text": "Contents",
                    "insertion_location": "start",
                    "replace_existing": True,
                },
                "last_status": "",
            },
        ]
        window._refresh_configured_pipeline()
        window.run_pipeline_from_ui()

        deadline = time.time() + 20
        while window.pipeline_worker is not None and time.time() < deadline:
            app.processEvents()
            time.sleep(0.05)
        assert window.pipeline_worker is None
        assert not output_path.exists()
        assert window.last_run_summary is not None
        assert window.last_run_summary["ok_count"] == 1
        assert window.last_run_summary["error_count"] == 1
        assert window.last_run_summary["saved"] is False
        assert window.configured_operations[0]["last_status"] == "OK"
        assert window.configured_operations[1]["last_status"] == "ERROR"
        assert any(entry["category"] == "operation" for entry in window.session_log_entries)
        assert any(
            entry["level"] == "error" and entry["category"] == "operation"
            for entry in window.session_log_entries
        )

        saved_log = window.save_session_log(str(log_path))
        assert saved_log == str(log_path)
        assert log_path.is_file()
        log_text = log_path.read_text(encoding="utf-8")
        assert '"category": "operation"' in log_text
        assert '"level": "error"' in log_text

        print(
            "Phase 9 smoke test passed: structured session logging, operation-vs-engine "
            "failure classification, and end-of-run summary data were produced and saved."
        )

    window.close()
    app.quit()
    return 0


def run_phase10_smoke() -> int:
    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase10_") as temp_dir:
        temp_root = Path(temp_dir)
        config_path = temp_root / "phase10-config.json"
        input_path = temp_root / "phase10-input.docx"
        output_dir = temp_root / "outputs"
        log_path = temp_root / "logs" / "phase10-session.jsonl"
        create_phase5_sample_doc(input_path, include_header=False)

        window = MainWindow(
            app_title=APP_TITLE,
            show_summary_dialog=False,
            config_path=config_path,
        )
        window.config_tab.engine_preference_combo.setCurrentIndex(
            window.config_tab.engine_preference_combo.findData("docx")
        )
        window.config_tab.error_policy_combo.setCurrentIndex(
            window.config_tab.error_policy_combo.findData("continue_on_error")
        )
        window.config_tab.log_verbosity_combo.setCurrentIndex(
            window.config_tab.log_verbosity_combo.findData("detailed")
        )
        window.config_tab.default_output_row.set_text(str(output_dir))
        window.config_tab.log_file_row.set_text(str(log_path))
        window.load_document(str(input_path), str(output_dir / "phase10-output.docx"))
        window.configured_operations = [
            {
                "display_name": "Find/Replace",
                "params": {
                    "find_text": "مرحبا",
                    "replace_text": "أهلا",
                    "scope": {"body": True},
                    "match_case": True,
                    "whole_word": False,
                },
                "last_status": "",
            },
            {
                "display_name": "Arabic RTL Normalize",
                "params": {
                    "scope": {"body": True},
                    "alignment": "right",
                    "direction": "rtl",
                    "font_name": "Arial",
                    "arabic_only": True,
                    "normalize_tables": True,
                    "normalize_lists": True,
                },
                "last_status": "",
            },
        ]
        window._refresh_configured_pipeline()
        window.save_pipeline_preset_named("Arabic cleanup")
        window.save_config_from_ui()
        assert config_path.is_file()
        window.close()

        reopened = MainWindow(
            app_title=APP_TITLE,
            show_summary_dialog=False,
            config_path=config_path,
        )
        assert reopened.config_tab.engine_preference_combo.currentData() == "docx"
        assert reopened.config_tab.error_policy_combo.currentData() == "continue_on_error"
        assert reopened.config_tab.log_verbosity_combo.currentData() == "detailed"
        assert reopened.config_tab.default_output_row.text() == str(output_dir)
        assert reopened.config_tab.log_file_row.text() == str(log_path)
        assert reopened.config_tab.last_input_value.text() == str(input_path)
        assert reopened.config_tab.presets_list.count() == 1
        assert reopened.config_tab.presets_list.item(0).text() == "Arabic cleanup"

        reopened.config_tab.presets_list.setCurrentRow(0)
        reopened.load_selected_preset()
        assert len(reopened.configured_operations) == 2
        assert reopened.configured_operations[0]["display_name"] == "Find/Replace"
        assert reopened.configured_operations[1]["display_name"] == "Arabic RTL Normalize"

        reopened.delete_selected_preset()
        assert reopened.config_tab.presets_list.count() == 0
        reopened.close()

        print(
            "Phase 10 smoke test passed: config settings, logging defaults, and "
            "pipeline presets persisted across restart and remained editable."
        )

    app.quit()
    return 0


def create_phase11_sample_doc(output_path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
    except Exception as exc:  # noqa: BLE001
        raise EngineUnavailableError(
            "python-docx is required to build the Phase 11 smoke sample."
        ) from exc

    document = Document()
    if "Custom Body" not in [style.name for style in document.styles]:
        document.styles.add_style("Custom Body", WD_STYLE_TYPE.PARAGRAPH)
    document.add_heading("Main Heading", level=1)
    document.add_heading("Sub Heading", level=2)
    document.add_paragraph("Body text", style="Custom Body")
    document.save(str(output_path))


def run_phase11_smoke() -> int:
    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase11_") as temp_dir:
        temp_root = Path(temp_dir)
        config_path = temp_root / "phase11-config.json"
        input_path = temp_root / "phase11-input.docx"
        output_path = temp_root / "phase11-output.docx"
        create_phase11_sample_doc(input_path)

        window = MainWindow(
            app_title=APP_TITLE,
            show_summary_dialog=False,
            config_path=config_path,
        )
        window.config_tab.engine_preference_combo.setCurrentIndex(
            window.config_tab.engine_preference_combo.findData("docx")
        )
        window.load_document(str(input_path), str(output_path))

        assert window.current_style_catalog["type_counts"].get("paragraph", 0) >= 3
        assert "Paragraph styles:" in window.pipeline_tab.styles_detail.text()
        assert "Recommended TOC paragraph styles:" in window.pipeline_tab.toc_style_hint.text()

        window.pipeline_tab.style_type_filter.setCurrentIndex(
            window.pipeline_tab.style_type_filter.findData("paragraph")
        )
        app.processEvents()
        paragraph_rows = [
            window.pipeline_tab.styles_list.item(index).text()
            for index in range(window.pipeline_tab.styles_list.count())
        ]
        assert paragraph_rows
        assert all("[paragraph]" in row for row in paragraph_rows)

        window.pipeline_tab.style_search_edit.setText("Heading")
        app.processEvents()
        filtered_rows = [
            window.pipeline_tab.styles_list.item(index).text()
            for index in range(window.pipeline_tab.styles_list.count())
        ]
        assert filtered_rows
        assert all("heading" in row.lower() for row in filtered_rows)

        toc_template = window._build_operation_template("Custom TOC Builder")
        assert toc_template["style_levels"]
        assert toc_template["style_levels"][0]["style_name"].startswith("Heading")

        print(
            "Phase 11 smoke test passed: style discovery supports typed filtering, "
            "clear summaries for odd documents, and TOC-oriented paragraph style suggestions."
        )

        window.close()

    app.quit()
    return 0


def create_phase12_basic_doc(output_path: Path, lines: list[str]) -> None:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise EngineUnavailableError(
            "python-docx is required to build the Phase 12 smoke sample."
        ) from exc

    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(str(output_path))


def wait_for_pipeline(window: MainWindow, app, timeout_seconds: float = 20.0) -> None:
    deadline = time.time() + timeout_seconds
    while window.pipeline_worker is not None and time.time() < deadline:
        app.processEvents()
        time.sleep(0.05)
    assert window.pipeline_worker is None


def run_phase12_smoke() -> int:
    from docx import Document

    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)

    with tempfile.TemporaryDirectory(prefix="batch_edit_phase12_") as temp_dir:
        temp_root = Path(temp_dir)
        work_dir = temp_root / "phase 12 عربي docs"
        work_dir.mkdir(parents=True, exist_ok=True)
        config_path = temp_root / "phase12-config.json"

        source_path = work_dir / "source ملف.docx"
        first_output = work_dir / "output one.docx"
        second_output = work_dir / "output two.docx"
        empty_source = work_dir / "empty doc.docx"

        create_phase12_basic_doc(source_path, ["مرحبا", "alpha beta"])
        create_phase12_basic_doc(empty_source, [])

        window = MainWindow(
            app_title=APP_TITLE,
            show_summary_dialog=False,
            config_path=config_path,
        )
        window.config_tab.engine_preference_combo.setCurrentIndex(
            window.config_tab.engine_preference_combo.findData("docx")
        )
        window.config_tab.error_policy_combo.setCurrentIndex(
            window.config_tab.error_policy_combo.findData("continue_on_error")
        )
        window.config_tab.log_verbosity_combo.setCurrentIndex(
            window.config_tab.log_verbosity_combo.findData("detailed")
        )
        window.config_tab.default_output_row.set_text(str(work_dir))

        # Repeated real-document runs in a unicode/spaced path should remain stable.
        window.load_document(str(source_path), str(first_output))
        window.configured_operations = [
            {
                "display_name": "Find/Replace",
                "params": {
                    "find_text": "مرحبا",
                    "replace_text": "أهلا",
                    "scope": {"body": True},
                    "match_case": True,
                    "whole_word": False,
                },
                "last_status": "",
            },
            {
                "display_name": "Arabic RTL Normalize",
                "params": {
                    "scope": {"body": True},
                    "alignment": "right",
                    "direction": "rtl",
                    "font_name": "Arial",
                    "arabic_only": True,
                    "normalize_tables": True,
                    "normalize_lists": True,
                },
                "last_status": "",
            },
        ]
        window._refresh_configured_pipeline()
        window.run_pipeline_from_ui()
        wait_for_pipeline(window, app)
        assert first_output.is_file()
        first_doc = Document(str(first_output))
        assert any("أهلا" in paragraph.text for paragraph in first_doc.paragraphs)

        # A second run on the same window should not leak stale state.
        window.load_document(str(first_output), str(second_output))
        window.configured_operations = [
            {
                "display_name": "Find/Replace",
                "params": {
                    "find_text": "alpha",
                    "replace_text": "omega",
                    "scope": {"body": True},
                    "match_case": True,
                    "whole_word": False,
                },
                "last_status": "",
            }
        ]
        window._refresh_configured_pipeline()
        window.run_pipeline_from_ui()
        wait_for_pipeline(window, app)
        assert second_output.is_file()
        second_doc = Document(str(second_output))
        assert any("omega beta" in paragraph.text for paragraph in second_doc.paragraphs)

        # Continue-on-error should save successful later operations and report both outcomes.
        mixed_output = work_dir / "mixed results.docx"
        window.load_document(str(source_path), str(mixed_output))
        window.configured_operations = [
            {
                "display_name": "Custom TOC Builder",
                "params": {
                    "style_levels": [{"style_name": "Heading 1", "level": 1}],
                    "title_text": "Contents",
                    "insertion_location": "start",
                    "replace_existing": True,
                },
                "last_status": "",
            },
            {
                "display_name": "Find/Replace",
                "params": {
                    "find_text": "alpha",
                    "replace_text": "gamma",
                    "scope": {"body": True},
                    "match_case": True,
                    "whole_word": False,
                },
                "last_status": "",
            },
        ]
        window._refresh_configured_pipeline()
        window.run_pipeline_from_ui()
        wait_for_pipeline(window, app)
        assert mixed_output.is_file()
        mixed_doc = Document(str(mixed_output))
        assert any("gamma beta" in paragraph.text for paragraph in mixed_doc.paragraphs)
        assert window.last_run_summary is not None
        assert window.last_run_summary["ok_count"] == 1
        assert window.last_run_summary["error_count"] == 1
        assert window.last_run_summary["saved"] is True

        # Odd-document loading should remain safe and informative.
        window.load_document(str(empty_source), str(work_dir / "empty output.docx"))
        assert "Built-in only document:" in window.pipeline_tab.styles_detail.text()
        window.clear_loaded_document()
        assert window.pipeline_tab.styles_list.count() == 0
        assert window.pipeline_tab.operations_list.count() == len(
            OperationRegistry.list_operations()
        )
        assert window.current_engine is None

        print(
            "Phase 12 smoke test passed: repeated real-document runs, continue-on-error "
            "handling, unicode/spaced paths, and odd-document load/reset behavior remained stable."
        )

        window.close()

    app.quit()
    return 0


def run_gui() -> int:
    app = QtWidgets.QApplication.instance() or QtWidgets.QApplication(sys.argv)
    window = MainWindow(app_title=APP_TITLE)
    window.show()
    return app.exec_()


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Batch Edit Files desktop app")
    parser.add_argument(
        "--smoke",
        choices=["phase1", "phase2", "phase3", "phase4", "phase5", "phase6", "phase7", "phase8", "phase9", "phase10", "phase11", "phase12"],
        help="Run a non-interactive smoke test for a completed phase.",
    )
    args = parser.parse_args(argv)

    if args.smoke:
        os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

    try:
        if args.smoke == "phase1":
            return run_phase1_smoke()
        if args.smoke == "phase2":
            return run_phase2_smoke()
        if args.smoke == "phase3":
            return run_phase3_smoke()
        if args.smoke == "phase4":
            return run_phase4_smoke()
        if args.smoke == "phase5":
            return run_phase5_smoke()
        if args.smoke == "phase6":
            return run_phase6_smoke()
        if args.smoke == "phase7":
            return run_phase7_smoke()
        if args.smoke == "phase8":
            return run_phase8_smoke()
        if args.smoke == "phase9":
            return run_phase9_smoke()
        if args.smoke == "phase10":
            return run_phase10_smoke()
        if args.smoke == "phase11":
            return run_phase11_smoke()
        if args.smoke == "phase12":
            return run_phase12_smoke()
        return run_gui()
    except (EngineUnavailableError, FileNotFoundError) as exc:
        print(str(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
