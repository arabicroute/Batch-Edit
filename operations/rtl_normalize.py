from __future__ import annotations

import re
from dataclasses import dataclass

from core.engines import ComEngine, DocxEngine
from core.operations import (
    Operation,
    OperationRegistry,
    OperationResult,
    OperationValidationError,
)

try:
    from docx.enum.table import WD_TABLE_DIRECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:  # noqa: BLE001
    WD_TABLE_DIRECTION = None  # type: ignore
    WD_ALIGN_PARAGRAPH = None  # type: ignore
    OxmlElement = None  # type: ignore
    qn = None  # type: ignore


ARABIC_TEXT_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")
COM_ALIGN_LEFT = 0
COM_ALIGN_CENTER = 1
COM_ALIGN_RIGHT = 2
COM_READING_ORDER_LTR = 1
COM_READING_ORDER_RTL = 2


@dataclass(slots=True)
class RtlScope:
    body: bool = True
    comments: bool = False
    headers_footers: bool = False

    @classmethod
    def from_params(cls, params: dict) -> "RtlScope":
        raw_scope = params.get("scope")
        if isinstance(raw_scope, dict):
            return cls(
                body=bool(raw_scope.get("body", True)),
                comments=bool(raw_scope.get("comments", False)),
                headers_footers=bool(raw_scope.get("headers_footers", False)),
            )
        return cls()

    def any_selected(self) -> bool:
        return self.body or self.comments or self.headers_footers

    def describe(self) -> str:
        parts: list[str] = []
        if self.body:
            parts.append("body")
        if self.comments:
            parts.append("comments")
        if self.headers_footers:
            parts.append("headers/footers")
        return ", ".join(parts)


def contains_arabic_text(value: str) -> bool:
    return bool(ARABIC_TEXT_RE.search(value or ""))


def ensure_on_off_child(parent, tag_name: str):
    element = parent.find(qn(tag_name))
    if element is None:
        element = OxmlElement(tag_name)
        parent.insert(0, element)
    element.set(qn("w:val"), "1")
    return element


def ensure_justification(p_pr, alignment_name: str) -> None:
    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), alignment_name)


def apply_run_rtl(run, font_name: str | None) -> None:
    r_pr = run._r.get_or_add_rPr()
    ensure_on_off_child(r_pr, "w:rtl")
    ensure_on_off_child(r_pr, "w:cs")
    if font_name:
        run.font.name = font_name
        if r_pr.find(qn("w:rFonts")) is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        else:
            r_fonts = r_pr.find(qn("w:rFonts"))
        for attr in ("ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{attr}"), font_name)


def apply_docx_paragraph_rtl(paragraph, alignment_name: str, font_name: str | None) -> bool:
    text = paragraph.text or ""
    p_pr = paragraph._p.get_or_add_pPr()
    ensure_on_off_child(p_pr, "w:bidi")
    ensure_justification(p_pr, alignment_name)
    if WD_ALIGN_PARAGRAPH is not None:
        alignment_map = {
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
        }
        paragraph.alignment = alignment_map.get(alignment_name, WD_ALIGN_PARAGRAPH.RIGHT)
    for run in paragraph.runs:
        if run.text:
            apply_run_rtl(run, font_name)
    return bool(text.strip())


def apply_docx_table_rtl(table) -> int:
    if WD_TABLE_DIRECTION is not None:
        table.direction = WD_TABLE_DIRECTION.RTL
    table._tbl.bidiVisual_val = True
    updated = 0
    for row in table.rows:
        for cell in row.cells:
            updated += 1
    return updated


def iter_docx_table_paragraphs(document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield table, paragraph


def iter_com_scope_ranges(document, scope: RtlScope):
    if scope.body:
        yield "Body", document.Content

    if scope.comments:
        comments = getattr(document, "Comments", None)
        if comments is not None:
            for index in range(1, comments.Count + 1):
                try:
                    comment = comments.Item(index)
                    yield f"Comment#{index}", comment.Range
                except Exception:  # noqa: BLE001
                    continue

    if scope.headers_footers:
        sections = getattr(document, "Sections", None)
        if sections is None:
            return
        for section_index in range(1, sections.Count + 1):
            section = sections.Item(section_index)
            headers = getattr(section, "Headers", None)
            footers = getattr(section, "Footers", None)
            if headers is not None:
                for item_index in range(1, headers.Count + 1):
                    try:
                        header = headers.Item(item_index)
                        yield f"Header S{section_index}#{item_index}", header.Range
                    except Exception:  # noqa: BLE001
                        continue
            if footers is not None:
                for item_index in range(1, footers.Count + 1):
                    try:
                        footer = footers.Item(item_index)
                        yield f"Footer S{section_index}#{item_index}", footer.Range
                    except Exception:  # noqa: BLE001
                        continue


@OperationRegistry.register("Arabic RTL Normalize")
class ArabicRtlNormalizeOperation(Operation):
    name = "Arabic RTL Normalize"
    required_capabilities = {"rtl_basic"}
    params_schema = {
        "scope": {"type": "object"},
        "alignment": {"type": "string", "default": "right"},
        "direction": {"type": "string", "default": "rtl"},
        "font_name": {"type": "string", "nullable": True},
        "arabic_only": {"type": "boolean", "default": True},
        "normalize_tables": {"type": "boolean", "default": True},
        "normalize_lists": {"type": "boolean", "default": True},
    }

    def validate(self, engine) -> None:
        self.ensure_supported(engine)
        scope = RtlScope.from_params(self.params)
        if not scope.any_selected():
            raise OperationValidationError("At least one scope must be selected.")

        alignment = str(self.params.get("alignment", "right")).lower()
        if alignment not in {"right", "center", "left"}:
            raise OperationValidationError(
                "Alignment must be one of: right, center, left."
            )

        direction = str(self.params.get("direction", "rtl")).lower()
        if direction != "rtl":
            raise OperationValidationError("This operation currently supports RTL only.")

        if isinstance(engine, DocxEngine):
            if scope.comments or scope.headers_footers:
                raise OperationValidationError(
                    "python-docx fallback supports body-only RTL normalization."
                )
        else:
            if scope.comments and "comments" not in engine.capabilities:
                raise OperationValidationError(
                    "The current engine does not support comment RTL normalization."
                )
            if scope.headers_footers and "headers_footers" not in engine.capabilities:
                raise OperationValidationError(
                    "The current engine does not support header/footer RTL normalization."
                )

    def describe(self) -> str:
        scope = RtlScope.from_params(self.params).describe()
        mode = "arabic-only" if self.params.get("arabic_only", True) else "force-all"
        return f"{self.name}: {scope} [{mode}]"

    def run(self, engine) -> OperationResult:
        self.validate(engine)
        if isinstance(engine, ComEngine):
            return self._run_com(engine)
        if isinstance(engine, DocxEngine):
            return self._run_docx(engine)
        raise OperationValidationError("Unsupported engine type for Arabic RTL Normalize.")

    def _should_touch(self, text: str) -> bool:
        if self.params.get("arabic_only", True):
            return contains_arabic_text(text)
        return bool((text or "").strip())

    def _run_docx(self, engine: DocxEngine) -> OperationResult:
        document = engine.document
        alignment_name = str(self.params.get("alignment", "right")).lower()
        font_name = self.params.get("font_name")
        normalize_tables = bool(self.params.get("normalize_tables", True))

        settings_element = document.settings.element
        ensure_on_off_child(settings_element, "w:bidi")

        paragraphs_updated = 0
        tables_updated = 0
        for paragraph in document.paragraphs:
            if not self._should_touch(paragraph.text):
                continue
            if apply_docx_paragraph_rtl(paragraph, alignment_name, font_name):
                paragraphs_updated += 1

        if normalize_tables:
            seen_tables: set[int] = set()
            for table, paragraph in iter_docx_table_paragraphs(document):
                if self._should_touch(paragraph.text):
                    apply_docx_paragraph_rtl(paragraph, alignment_name, font_name)
                table_id = id(table)
                if table_id not in seen_tables:
                    tables_updated += 1 if apply_docx_table_rtl(table) >= 0 else 0
                    seen_tables.add(table_id)

        return OperationResult(
            status="ok",
            message=(
                f"Normalized {paragraphs_updated} paragraph(s) and "
                f"{tables_updated} table(s) for RTL via {engine.engine_name}."
            ),
            details={
                "paragraphs_updated": paragraphs_updated,
                "tables_updated": tables_updated,
                "engine": engine.engine_name,
                "scope": "body",
            },
        )

    def _run_com(self, engine: ComEngine) -> OperationResult:
        document = engine.document
        scope = RtlScope.from_params(self.params)
        alignment_name = str(self.params.get("alignment", "right")).lower()
        font_name = self.params.get("font_name")
        normalize_tables = bool(self.params.get("normalize_tables", True))
        normalize_lists = bool(self.params.get("normalize_lists", True))

        alignment_map = {
            "right": COM_ALIGN_RIGHT,
            "center": COM_ALIGN_CENTER,
            "left": COM_ALIGN_LEFT,
        }
        paragraphs_updated = 0
        tables_updated = 0

        for _, rng in iter_com_scope_ranges(document, scope):
            paragraphs = getattr(rng, "Paragraphs", None)
            if paragraphs is not None:
                for index in range(1, paragraphs.Count + 1):
                    paragraph = paragraphs.Item(index)
                    text = str(paragraph.Range.Text).strip()
                    if not self._should_touch(text):
                        continue
                    fmt = paragraph.Range.ParagraphFormat
                    fmt.Alignment = alignment_map[alignment_name]
                    try:
                        fmt.ReadingOrder = COM_READING_ORDER_RTL
                    except Exception:  # noqa: BLE001
                        pass
                    if normalize_lists:
                        try:
                            fmt.MirrorIndents = True
                        except Exception:  # noqa: BLE001
                            pass
                    if font_name:
                        try:
                            paragraph.Range.Font.Name = font_name
                            paragraph.Range.Font.NameBi = font_name
                        except Exception:  # noqa: BLE001
                            paragraph.Range.Font.Name = font_name
                    paragraphs_updated += 1

            if normalize_tables:
                tables = getattr(rng, "Tables", None)
                if tables is not None:
                    for index in range(1, tables.Count + 1):
                        table = tables.Item(index)
                        try:
                            table.Rows.Alignment = alignment_map[alignment_name]
                        except Exception:  # noqa: BLE001
                            pass
                        try:
                            table.Range.ParagraphFormat.ReadingOrder = COM_READING_ORDER_RTL
                        except Exception:  # noqa: BLE001
                            pass
                        if font_name:
                            try:
                                table.Range.Font.Name = font_name
                                table.Range.Font.NameBi = font_name
                            except Exception:  # noqa: BLE001
                                table.Range.Font.Name = font_name
                        tables_updated += 1

        return OperationResult(
            status="ok",
            message=(
                f"Normalized {paragraphs_updated} paragraph(s) and "
                f"{tables_updated} table(s) for RTL via {engine.engine_name}."
            ),
            details={
                "paragraphs_updated": paragraphs_updated,
                "tables_updated": tables_updated,
                "engine": engine.engine_name,
                "scope": scope.describe(),
            },
        )
